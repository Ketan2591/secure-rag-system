from pathlib import Path
from typing import BinaryIO

import fitz
from docx import Document
from langchain_core.documents import Document as LangChainDocument


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def _get_extension(filename: str) -> str:
    """Return the lowercase file extension."""
    return Path(filename).suffix.lower()


def validate_file(filename: str) -> None:
    """Validate whether the uploaded file type is supported."""
    extension = _get_extension(filename)

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {extension or 'unknown'}. "
            "Only PDF, DOCX, and TXT files are supported."
        )


def process_pdf(file: BinaryIO, filename: str) -> list[LangChainDocument]:
    """Extract text from a PDF while preserving page numbers."""
    documents = []

    file.seek(0)
    pdf_bytes = file.read()

    with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf:
        for page_index, page in enumerate(pdf):
            text = page.get_text("text").strip()

            if not text:
                continue

            documents.append(
                LangChainDocument(
                    page_content=text,
                    metadata={
                        "source": filename,
                        "file_type": "pdf",
                        "page": page_index + 1,
                    },
                )
            )

    return documents


def process_docx(file: BinaryIO, filename: str) -> list[LangChainDocument]:
    """Extract text from a DOCX document."""
    file.seek(0)
    document = Document(file)

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    text = "\n".join(paragraphs)

    if not text:
        return []

    return [
        LangChainDocument(
            page_content=text,
            metadata={
                "source": filename,
                "file_type": "docx",
                "page": 1,
            },
        )
    ]


def process_txt(file: BinaryIO, filename: str) -> list[LangChainDocument]:
    """Extract text from a UTF-8 text file."""
    file.seek(0)
    raw_data = file.read()

    if isinstance(raw_data, bytes):
        try:
            text = raw_data.decode("utf-8")
        except UnicodeDecodeError:
            text = raw_data.decode("utf-8", errors="replace")
    else:
        text = raw_data

    text = text.strip()

    if not text:
        return []

    return [
        LangChainDocument(
            page_content=text,
            metadata={
                "source": filename,
                "file_type": "txt",
                "page": 1,
            },
        )
    ]


def process_document(
    file: BinaryIO,
    filename: str,
) -> list[LangChainDocument]:
    """
    Validate and process an uploaded document.

    Returns LangChain Documents containing extracted text and metadata.
    """
    validate_file(filename)

    extension = _get_extension(filename)

    if extension == ".pdf":
        documents = process_pdf(file, filename)

    elif extension == ".docx":
        documents = process_docx(file, filename)

    else:
        documents = process_txt(file, filename)

    if not documents:
        raise ValueError(
            f"No readable text was found in '{filename}'. "
            "The document may be empty or image-based."
        )

    return documents